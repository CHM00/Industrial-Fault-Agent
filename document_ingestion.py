"""Safe text extraction and deterministic chunking for managed knowledge."""

from __future__ import annotations

import csv
import io
import json
import os
import re
from pathlib import Path


MAX_UPLOAD_BYTES = int(os.environ.get("KNOWLEDGE_MAX_UPLOAD_BYTES", str(15 * 1024 * 1024)))
MAX_EXTRACTED_CHARS = int(os.environ.get("KNOWLEDGE_MAX_EXTRACTED_CHARS", "2000000"))


class DocumentIngestionError(ValueError):
    pass


def _decode(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentIngestionError("无法识别文本编码，请转换为 UTF-8")


def extract_text(filename: str, data: bytes) -> tuple[str, str]:
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentIngestionError(f"文件超过上传限制：{len(data)} > {MAX_UPLOAD_BYTES}")
    suffix = Path(filename or "document.txt").suffix.lower()
    try:
        if suffix in {".txt", ".md", ".markdown", ".log"}:
            text = _decode(data)
        elif suffix == ".json":
            value = json.loads(_decode(data))
            text = json.dumps(value, ensure_ascii=False, indent=2)
        elif suffix in {".csv", ".tsv"}:
            delimiter = "\t" if suffix == ".tsv" else ","
            rows = csv.reader(io.StringIO(_decode(data)), delimiter=delimiter)
            text = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
        elif suffix == ".docx":
            from docx import Document

            doc = Document(io.BytesIO(data))
            lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    lines.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n".join(lines)
        elif suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join(
                f"[第 {index + 1} 页]\n{page.extract_text() or ''}"
                for index, page in enumerate(reader.pages)
            )
        elif suffix in {".xlsx", ".xlsm"}:
            from openpyxl import load_workbook

            workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            lines = []
            for sheet in workbook.worksheets:
                lines.append(f"[工作表：{sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(value).strip() if value is not None else "" for value in row]
                    if any(cells):
                        lines.append(" | ".join(cells))
            text = "\n".join(lines)
        else:
            raise DocumentIngestionError(
                "不支持的文件类型；支持 txt/md/csv/tsv/json/pdf/docx/xlsx/xlsm"
            )
    except DocumentIngestionError:
        raise
    except Exception as exc:
        raise DocumentIngestionError(f"文档解析失败：{exc}") from exc

    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text).strip()
    if not text:
        raise DocumentIngestionError("文档未提取到可索引文本")
    if len(text) > MAX_EXTRACTED_CHARS:
        raise DocumentIngestionError(
            f"提取文本超过限制：{len(text)} > {MAX_EXTRACTED_CHARS}"
        )
    return text, suffix.lstrip(".") or "txt"


def chunk_text(text: str, max_chars: int = 900, overlap: int = 120) -> list[dict]:
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]
    chunks: list[dict] = []
    buffer = ""
    location = ""
    buffer_location = ""
    for paragraph in paragraphs:
        page = re.match(r"\[第\s*(\d+)\s*页\]", paragraph)
        sheet = re.match(r"\[工作表：(.+?)\]", paragraph)
        next_location = location
        if page:
            next_location = f"第 {page.group(1)} 页"
        elif sheet:
            next_location = f"工作表 {sheet.group(1)}"
        if buffer and len(buffer) + len(paragraph) + 2 > max_chars:
            chunks.append({"content": buffer, "location": buffer_location or f"分块 {len(chunks)+1}"})
            buffer = buffer[-overlap:] + "\n\n" + paragraph
            buffer_location = next_location or location
        else:
            buffer = f"{buffer}\n\n{paragraph}".strip()
            if not buffer_location:
                buffer_location = next_location or location
        location = next_location
        while len(buffer) > max_chars * 2:
            chunks.append({"content": buffer[:max_chars], "location": buffer_location or f"分块 {len(chunks)+1}"})
            buffer = buffer[max_chars - overlap :]
    if buffer:
        chunks.append({"content": buffer, "location": buffer_location or location or f"分块 {len(chunks)+1}"})
    return chunks
