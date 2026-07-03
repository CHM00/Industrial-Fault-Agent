"""
Convert a markdown file to a formatted Word document (`.docx`).
"""

import markdown
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------- constants ----------
MD_FILE = Path(r"d:\git_repo_file\LangGraph_Agent_Work\期末大作业说明文档.md")
DOCX_FILE = Path(r"d:\git_repo_file\LangGraph_Agent_Work\期末大作业.docx")

FONT_BODY = "Calibri"
FONT_CODE = "Courier New"
FONT_TITLE_CJK = "Calibri"
FONT_BODY_CJK = "宋体"
FONT_CODE_CJK = "Courier New"

HEADING_SIZES = {
    1: 16,
    2: 14,
    3: 12,
    4: 11,
    5: 10,
    6: 10,
}

BODY_SIZE = 11
LINE_SPACING = 1.15
PAGE_MARGIN = 2.5  # cm


def _set_font(run, size=None, bold=False, italic=False, font_name=None):
    """Apply font properties to a run."""
    r = run.font
    r.size = Pt(size) if size else r.size
    r.bold = bold
    r.italic = italic
    r.name = font_name or FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY_CJK)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name or FONT_BODY)


def _set_paragraph_spacing(p, space_before=6, space_after=6):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(BODY_SIZE * LINE_SPACING)


def _apply_table_style(table):
    """Apply clean formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = Pt(BODY_SIZE * 1.05)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(BODY_SIZE - 1)
                    run.font.name = FONT_BODY
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY_CJK)


def _set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(PAGE_MARGIN)
        section.bottom_margin = Cm(PAGE_MARGIN)
        section.left_margin = Cm(PAGE_MARGIN)
        section.right_margin = Cm(PAGE_MARGIN)


def _is_fenced_code_block(line):
    return line.strip().startswith("```")


def _build_doc_from_md(text: str) -> Document:
    doc = Document()
    _set_margins(doc)

    # Default style tweaks
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(BODY_SIZE)
    font.east_asian_font = FONT_BODY_CJK
    style.paragraph_format.line_spacing = Pt(BODY_SIZE * LINE_SPACING)
    style.paragraph_format.space_before = Pt(4)
    style.paragraph_format.space_after = Pt(4)

    lines = text.split("\n")
    i = 0
    in_code_fence = False
    code_lines = []
    code_lang = ""
    in_blockquote = False
    blockquote_lines = []
    in_list = False
    list_items = []
    list_type = None

    def _flush_list():
        nonlocal in_list, list_items, list_type
        if not list_items:
            return
        prefix = "" if list_type == "ordered" else ""
        for idx, item in enumerate(list_items, 1):
            p = doc.add_paragraph(style="List Bullet") if list_type != "ordered" else doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = Pt(BODY_SIZE * 1.05)
            if list_type == "ordered":
                p.paragraph_format.left_indent = Cm(1.27)
                txt = f"{idx}. {item}"
                run = p.add_run(txt)
            else:
                run = p.add_run(item)
            _set_font(run, size=BODY_SIZE)
        in_list = False
        list_items.clear()
        list_type = None

    def _flush_blockquote():
        nonlocal in_blockquote, blockquote_lines
        if not blockquote_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(BODY_SIZE * 1.05)
        # strip leading >
        clean = "\n".join(
            re.sub(r"^\s*>\s?", "", line) for line in blockquote_lines
        )
        run = p.add_run(clean)
        _set_font(run, size=BODY_SIZE, italic=True)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        in_blockquote = False
        blockquote_lines.clear()

    def _flush_code():
        nonlocal in_code_fence, code_lines, code_lang
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(12)
        lang_label = code_lang if code_lang else ""
        if lang_label:
            run = p.add_run(f"[{lang_label}]\n")
            _set_font(run, size=9, bold=True)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for line in code_lines:
            run = p.add_run(line + "\n")
            _set_font(run, size=9, font_name=FONT_CODE)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE_CJK)
        p._element.get_or_add_pPr().append(
            parse_xml(
                '<w:shd {} w:fill="F5F5F5" w:val="clear"/>'.format(nsdecls("w"))
            )
        )
        in_code_fence = False
        code_lines.clear()
        code_lang = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if _is_fenced_code_block(stripped):
            if not in_code_fence:
                _flush_blockquote()
                _flush_list()
                in_code_fence = True
                lang_match = re.match(r"```(\w*)", stripped)
                code_lang = lang_match.group(1) if lang_match else ""
            else:
                _flush_code()
            i += 1
            continue

        if in_code_fence:
            code_lines.append(line)
            i += 1
            continue

        # Blank line
        if not stripped:
            _flush_blockquote()
            _flush_list()
            _flush_code()
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            _flush_list()
            _flush_code()
            in_blockquote = True
            blockquote_lines.append(stripped)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*[-*_]\s*[-*_]\s*[-*_]", stripped):
            _flush_blockquote()
            _flush_list()
            _flush_code()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                '<w:pBdr {}>'.format(nsdecls("w"))
                + '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="AAAAAA"/>'
                + "</w:pBdr>"
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if heading_match:
            _flush_blockquote()
            _flush_list()
            _flush_code()
            level = len(heading_match.group(1))
            title_text = heading_match.group(2).strip()

            # Strip inline formatting markers for clean heading text
            clean_title = re.sub(r"\*\*([^*]+)\*\*", r"\1", title_text)
            clean_title = re.sub(r"\*([^*]+)\*", r"\1", clean_title)

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14 - level * 1)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = Pt(18 - level)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            bold = level <= 3
            sz = HEADING_SIZES.get(level, 11)

            run = p.add_run(clean_title)
            _set_font(run, size=sz, bold=bold)

            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            elif level == 2:
                run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
            elif level == 3:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            i += 1
            continue

        # Table detection
        if "|" in stripped and "^[-|]" in "\n".join(
            lines[i : i + 3]
        ):
            _flush_blockquote()
            _flush_list()
            _flush_code()
            # Gather rows until a non-table line
            rows = []
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line.startswith("|"):
                    break
                rows.append(row_line)
                i += 1
                if row_line and not row_line.startswith("|"):
                    break
                # Stop at separator row
                if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", row_line):
                    break
            if len(rows) >= 2:
                _insert_table(doc, rows)
            continue

        # Unordered list
        ul_match = re.match(r"^\s*[-*+]\s+(.+)", stripped)
        if ul_match:
            _flush_blockquote()
            _flush_code()
            in_list = True
            list_type = "unordered"
            list_items.append(ul_match.group(1).strip())
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^\s*\d+\.\s+(.+)", stripped)
        if ol_match:
            _flush_blockquote()
            _flush_code()
            in_list = True
            list_type = "ordered"
            list_items.append(ol_match.group(1).strip())
            i += 1
            continue

        # Regular paragraph
        _flush_blockquote()
        _flush_list()
        _flush_code()
        p = doc.add_paragraph()
        _set_paragraph_spacing(p)
        _add_inline_formatting(p, stripped)
        i += 1

    # Flush remaining
    _flush_blockquote()
    _flush_list()
    _flush_code()

    return doc


def _insert_table(doc, rows_raw):
    """Insert a parsed markdown table into the document."""
    # Remove separator rows
    clean_rows = []
    for row in rows_raw:
        if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", row):
            continue
        # strip trailing |
        clean = row.rstrip("|").strip()
        cells = [c.strip() for c in clean.split("|")]
        cells = [c for c in cells if c]
        if cells:
            clean_rows.append(cells)

    if len(clean_rows) < 1:
        return

    num_cols = max(len(r) for r in clean_rows)
    table = doc.add_table(rows=len(clean_rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row_data in enumerate(clean_rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _set_paragraph_spacing(p, space_before=2, space_after=2)
                run = p.add_run(cell_text)
                is_header = r_idx == 0
                _set_font(run, size=BODY_SIZE - 1, bold=is_header)

    _apply_table_style(table)


def _add_inline_formatting(paragraph, text: str):
    """Add bold/italic/code formatting to paragraph runs."""
    parts = []
    # We do a simple left-to-right scan picking out **bold**, *italic*, and `code`
    i = 0
    chars = list(text)
    while i < len(chars):
        # Bold
        if i < len(chars) - 1 and chars[i] == "*" and chars[i + 1] == "*":
            end = text.find("**", i + 2)
            if end == -1:
                end = i + 2
            else:
                end += 2
            segment = text[i + 2 : end - 2] if end > i + 2 else ""
            run = paragraph.add_run(segment)
            _set_font(run, size=BODY_SIZE, bold=True)
            i = end
        # Italic
        elif chars[i] == "*":
            end = text.find("*", i + 1)
            if end == -1:
                run = paragraph.add_run(chars[i])
                _set_font(run, size=BODY_SIZE)
                i += 1
            else:
                segment = text[i + 1 : end]
                run = paragraph.add_run(segment)
                _set_font(run, size=BODY_SIZE, italic=True)
                i = end + 1
        # Inline code
        elif chars[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                run = paragraph.add_run(chars[i])
                _set_font(run, size=BODY_SIZE)
                i += 1
            else:
                segment = text[i + 1 : end]
                run = paragraph.add_run(segment)
                _set_font(run, size=BODY_SIZE, font_name=FONT_CODE)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE_CJK)
                i = end + 1
        # Images ![alt](url)
        elif i < len(chars) - 1 and chars[i] == "!" and chars[i + 1] == "[":
            m = re.match(r"!\[([^\]]*)\]\([^)]*\)", text[i:])
            if m:
                alt = m.group(1)
                if alt:
                    run = paragraph.add_run(alt)
                    _set_font(run, size=BODY_SIZE, italic=True)
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                i += m.end()
            else:
                run = paragraph.add_run(chars[i])
                _set_font(run, size=BODY_SIZE)
                i += 1
        else:
            # Collect consecutive plain text characters
            start = i
            while i < len(chars) and chars[i] not in ("*", "`", "!"):
                i += 1
            if i > start:
                run = paragraph.add_run(text[start:i])
                _set_font(run, size=BODY_SIZE)
            else:
                # single special char not part of a pattern
                run = paragraph.add_run(chars[i])
                _set_font(run, size=BODY_SIZE)
                i += 1


def main():
    if not MD_FILE.exists():
        print(f"Error: markdown file not found: {MD_FILE}")
        sys.exit(1)

    md_text = MD_FILE.read_text(encoding="utf-8")
    doc = _build_doc_from_md(md_text)
    doc.save(str(DOCX_FILE))
    print(f"Saved -> {DOCX_FILE}")


if __name__ == "__main__":
    main()
