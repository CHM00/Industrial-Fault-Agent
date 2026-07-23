"""Generate Word, PDF and field checklist artifacts from a diagnosis state."""

from __future__ import annotations

import csv
import html
import io
import os
import re
from pathlib import Path


def _safe_name(value: str) -> str:
    return re.sub(r"[^\w\-]+", "_", value, flags=re.UNICODE)[:60] or "diagnosis"


def _sections(state: dict) -> list[tuple[str, str]]:
    safety = state.get("safety_assessment", {}) or {}
    controls = "\n".join(f"- {item}" for item in safety.get("controls", []))
    mappings = state.get("evidence_mappings", []) or []
    evidence_lines = []
    for mapping in mappings:
        evidence_lines.append(f"节点 {mapping.get('node_id')}：{mapping.get('label')}")
        for item in mapping.get("evidence", []):
            evidence_lines.append(
                f"  - {item.get('evidence_id')} {item.get('title')} "
                f"（{item.get('relation', 'supports')}，置信度 {item.get('confidence', 0):.2f}，{item.get('location', '')}）"
            )
    return [
        ("故障描述", state.get("fault_input", "")),
        ("设备与工况", state.get("structured_context", "") or "未提供"),
        ("安全预检", f"风险等级：{safety.get('risk_level', 'low')}\n{controls}"),
        ("内部知识", state.get("internal_knowledge", "")[:6000]),
        ("外部研究", state.get("external_knowledge", "")[:6000]),
        ("诊断上下文", state.get("filtered_context", "")[:8000]),
        ("审计结果", state.get("audit_result", "")),
        ("流程审批", state.get("procedure_context", "") or "未进入版本审批"),
        ("步骤证据", "\n".join(evidence_lines) or "暂无证据映射"),
        ("Mermaid 排障流程", state.get("mermaid_diagram", "")),
    ]


def export_docx(state: dict) -> tuple[bytes, str, str]:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading("工业故障诊断报告", level=0)
    document.add_paragraph(f"请求编号：{state.get('request_id', '')}")
    for title, content in _sections(state):
        document.add_heading(title, level=1)
        for line in str(content or "").splitlines():
            paragraph = document.add_paragraph(line)
            for run in paragraph.runs:
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(10.5)
    buffer = io.BytesIO()
    document.save(buffer)
    filename = f"{_safe_name(state.get('fault_input', 'diagnosis'))}.docx"
    return buffer.getvalue(), filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        os.environ.get("PDF_CJK_FONT", ""),
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).is_file():
            try:
                pdfmetrics.registerFont(TTFont("PilotCJK", candidate))
                return "PilotCJK"
            except Exception:
                continue
    return "Helvetica"


def export_pdf(state: dict) -> tuple[bytes, str, str]:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    font = _pdf_font()
    title_style = ParagraphStyle("TitleCJK", fontName=font, fontSize=18, leading=24, spaceAfter=10)
    heading_style = ParagraphStyle("HeadingCJK", fontName=font, fontSize=13, leading=18, spaceBefore=10, spaceAfter=6)
    body_style = ParagraphStyle("BodyCJK", fontName=font, fontSize=9.5, leading=14, alignment=TA_LEFT)
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
    )
    story = [Paragraph("工业故障诊断报告", title_style)]
    for title, content in _sections(state):
        story.append(Paragraph(html.escape(title), heading_style))
        lines = str(content or "").splitlines() or [""]
        for line in lines:
            story.append(Paragraph(html.escape(line) or "&nbsp;", body_style))
        story.append(Spacer(1, 4 * mm))
    doc.build(story)
    filename = f"{_safe_name(state.get('fault_input', 'diagnosis'))}.pdf"
    return buffer.getvalue(), filename, "application/pdf"


def export_checklist(state: dict) -> tuple[bytes, str, str]:
    mappings = state.get("evidence_mappings", []) or []
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["序号", "节点ID", "检查/操作步骤", "风险等级", "证据", "执行结果", "测量值", "执行人", "时间"])
    risk = (state.get("safety_assessment") or {}).get("risk_level", "low")
    for index, mapping in enumerate(mappings, 1):
        evidence = "; ".join(item.get("evidence_id", "") for item in mapping.get("evidence", []))
        writer.writerow([index, mapping.get("node_id", ""), mapping.get("label", ""), risk, evidence, "", "", "", ""])
    if not mappings:
        for index, line in enumerate(state.get("mermaid_diagram", "").splitlines()[1:], 1):
            writer.writerow([index, "", line.strip(), risk, "", "", "", "", ""])
    filename = f"{_safe_name(state.get('fault_input', 'diagnosis'))}_检查单.csv"
    return ("\ufeff" + output.getvalue()).encode("utf-8"), filename, "text/csv; charset=utf-8"


def export_artifact(format_name: str, state: dict) -> tuple[bytes, str, str]:
    if format_name == "docx":
        return export_docx(state)
    if format_name == "pdf":
        return export_pdf(state)
    if format_name in {"checklist", "csv"}:
        return export_checklist(state)
    raise ValueError("format 仅支持 docx、pdf、checklist")
